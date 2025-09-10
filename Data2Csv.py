from docx import Document
import csv
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell

def extract_question_and_code(input_file, cell):
    question_parts = []
    code_parts = []
    # ---------------- DOCX ----------------
    if input_file.endswith(".docx"):
        for para in cell.paragraphs:
            for run in para.runs:
                font = run.font
                if not (font.name == "Courier New"):
                    question_parts.append(run.text)

        for tbl in cell.tables:
            if len(tbl.rows) == 1 and len(tbl.columns) == 1:
                code_parts.append(tbl.rows[0].cells[0].text)
                code_parts.append("\n")

    # ---------------- XLSX ----------------
    if input_file.endswith(".xlsx"):
        if cell.value:
            # Font check (only works if workbook was not loaded with data_only=True)
            if cell.font.name == "Courier New":
                code_parts.append(str(cell.value))
            else:
                question_parts.append(str(cell.value))

    question_text = " ".join(q.strip() for q in question_parts if q.strip())
    code_text = "\n".join(c.strip() for c in code_parts if c.strip())

    return question_text, code_text


def document_to_data(input_file, output_file):
    # ---------------- DOCX ----------------
    structured_data = []
    index_counter = 1
    correct = []
    if input_file.endswith(".docx"):
        doc = Document(input_file)
        for table in doc.tables:
            for row_counter in range(1, len(table.rows)):
                #extract question and code:
                question_text, code_text = extract_question_and_code(input_file, table.rows[row_counter].cells[1])

                #extract answer
                answers_raw = table.rows[row_counter].cells[2].text.strip().split("\n")
                answers = [
                    a.strip()[3:].strip() if a.strip()[:2] in ["A.", "B.", "C.", "D."] else a.strip()
                    for a in answers_raw if a.strip()
                ]
                while len(answers) < 4:
                    answers.append("")

                #extract correct answer
                correct.append(table.rows[row_counter].cells[3].text.strip().lower())

                record = [
                    index_counter,   # index
                    "",              # context
                    question_text,   # question
                    "",              # image
                    "",              # audio
                    code_text,       # code
                    answers[0],      # A
                    answers[1],      # B
                    answers[2],      # C
                    answers[3],      # D
                    correct[row_counter-1],         # correct
                    "",              # hint
                    "",              # set_question_id
                    "",              # tags
                ]
                structured_data.append(record)
                index_counter += 1

    # ---------------- XLSX ----------------
    if input_file.endswith(".xlsx"):
        wb = load_workbook(input_file)
        sheet = wb.active
        for index, row in enumerate(sheet.iter_rows(min_row=2, values_only=False), start = 1):
            #extract question
            question_text, code_text = extract_question_and_code(input_file, row[1])

            # extract answer
            answers = []
            for i in range(2, 9, 2):
                answers.append(row[i].value.strip())
            while len(answers) < 4:
                answers.append("")

            #extract correct answer
            correct.append(row[10].value.strip())

            record = [
                index_counter,  # index
                "",  # context
                question_text,  # question
                "",  # image
                "",  # audio
                code_text,  # code
                answers[0],  # A
                answers[1],  # B
                answers[2],  # C
                answers[3],  # D
                correct[index-1],  # correct
                "",  # hint
                "",  # set_question_id
                "",  # tags
            ]
            structured_data.append(record)
            index_counter += 1

    headers = ["index", "context", "question", "image", "audio", "code",
               "A", "B", "C", "D", "correct", "hint", "set_question_id", "tags"]

    with open(output_file, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile, quoting=csv.QUOTE_MINIMAL)
        writer.writerow(headers)
        writer.writerows(structured_data)

document_to_data(
    r"E:/Document/NEU/ScoreUp 2.0/Word collection/Input/Cau hoi midterm/Cau hoi midterm.xlsx",
    r"E:/Document/NEU/ScoreUp 2.0/Word collection/Output/CSV/filtered_Cau hoi midterm_questions.csv"
)

print("Successful")