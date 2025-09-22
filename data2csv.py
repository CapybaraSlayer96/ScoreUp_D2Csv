from docx import Document
import csv
from openpyxl import load_workbook
import re

def extract_question_and_code(input_file, cell):
    question_parts = []
    code_parts = []
    # ---------------- DOCX ----------------
    if input_file.endswith(".docx"):
        for para in cell.paragraphs:
            for run in para.runs:
                font = run.font
                if not (font.name == "Courier New"):
                    question_parts.append(run.text)

        for tbl in cell.tables:
            if len(tbl.rows) == 1 and len(tbl.columns) == 1:
                code_parts.append(tbl.rows[0].cells[0].text)
                code_parts.append("\n")
        question_text = " ".join(q.strip() for q in question_parts if q.strip())
        code_text = "\n".join(c.strip() for c in code_parts if c.strip())

    # # ---------------- XLSX ----------------
    # if input_file.endswith(".xlsx"):
    #     # Regex tìm code block
    #     pattern = r"```(.*?)```"
    #     if cell.value:
    #         code_blocks = re.findall(pattern, cell.value, flags=re.DOTALL)
    #     else:
    #         code_blocks = "", ""
    #
    #     # Text thường (ngoài code)
    #     question_text = re.sub(pattern, "", cell.value, flags=re.DOTALL).strip()
    #
    #     # Làm sạch code block, giữ nguyên indent
    #     code_parts = [c.strip("\n") for c in code_blocks]
    #     code_text = "\n".join(c.rstrip() for c in code_parts if c.strip())

    return question_text, code_text

def safe_strip(val):
    return val.strip() if isinstance(val, str) else (val if val else "")

def document_to_data(input_file, output_file, tags, start_index):
    structured_data = []
    index_counter = start_index


    # ---------------- DOCX ----------------
    if input_file.endswith(".docx"):
        doc = Document(input_file)
        for table in doc.tables:
            for row_counter in range(1, len(table.rows)):
                #extract question and code:
                question_text, code_text = extract_question_and_code(input_file, table.rows[row_counter].cells[1])

                #extract answer
                cell_val = table.rows[row_counter].cells[2].text
                raw_answers = cell_val.strip().split("\n") if cell_val else []
                answers = []
                for a in raw_answers:
                    cleaned_answers = a.strip()
                    cleaned_answers = re.sub(r"^[A-D](?:[\.\)]|\s)\s*", "", cleaned_answers)
                    answers.append(cleaned_answers)
                while len(answers) < 4:
                    answers.append("")

                #extract correct answer
                correct = table.rows[row_counter].cells[3].text.strip().lower()

                record = [
                    index_counter,   # index
                    "",              # context
                    question_text,   # question
                    "",              # image
                    "",              # audio
                    code_text,       # code
                    answers[0],      # A
                    answers[1],      # B
                    answers[2],      # C
                    answers[3],      # D
                    correct,         # correct
                    "",              # hint
                    "",              # set_question_id
                    tags,             # tags
                ]
                structured_data.append(record)
                index_counter += 1

    # ---------------- XLSX ----------------
    if input_file.endswith(".xlsx"):
        wb = load_workbook(input_file)
        sheet = wb.active
        for row in sheet.iter_rows(min_row=2, values_only=False):
            #extract question
            question_text = safe_strip(row[0].value)

            #extract code
            code_text = ""
            code_raw = row[1].value
            if code_raw:
                code_parts = code_raw.splitlines()
                code_text = "\n".join(line.rstrip() for line in code_parts)

            # extract answer
            answers = []
            for i in range (2,6):
                answers.append(safe_strip(row[i].value))
            while len(answers) < 4:
                answers.append("")

            #extract correct answer
            correct = safe_strip(row[6].value)

            #extract hints
            # check = "Đúng,"
            # hint = ""
            # for i in range(2, 9, 2):
            #     cell_val = row[i].value
            #     if cell_val and isinstance(cell_val, str):
            #         stripped = cell_val.strip().lstrip(check).strip()
            #         if stripped != cell_val.strip():
            #             hint = stripped
            #             break
            hint = safe_strip(row[7].value)

            record = [
                index_counter,  # index
                "",             # context
                question_text,  # question
                "",             # image
                "",             # audio
                code_text,      # code
                answers[0],     # A
                answers[1],     # B
                answers[2],     # C
                answers[3],     # D
                correct,        # correct
                hint,           # hint
                "",             # set_question_id
                tags,            # tags
            ]
            structured_data.append(record)
            index_counter += 1

    headers = ["index", "context", "question", "image", "audio", "code",
               "A", "B", "C", "D", "correct", "hint", "set_question_id", "tags"]

    with open(output_file, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile, quoting=csv.QUOTE_MINIMAL)
        writer.writerow(headers)
        writer.writerows(structured_data)

#Testing
# document_to_data('E:/Document/NEU/Project/ScoreUp 2.0/Documents collection/Input/Chapter5.xlsx', 'E:/Document/NEU/Project/ScoreUp 2.0/Documents collection/Output/CSV/Chapter5_question.csv')
# print("Successful")